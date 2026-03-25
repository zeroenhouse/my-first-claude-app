#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import os

# docxファイルから売上データを抽出
doc = Document('GiGO我孫子_買収計画_総合まとめ.docx')

# テーブル7から売上試算データを抽出
sales_data = []
if len(doc.tables) > 7:
    table = doc.tables[7]
    for row_idx, row in enumerate(table.rows):
        if row_idx > 0:
            row_data = [cell.text for cell in row.cells]
            sale_scenario = row_data[0]
            annual_sales = row_data[3]
            sales_data.append({'scenario': sale_scenario, 'sales': annual_sales})
            print(f"売上データ: {sale_scenario} = {annual_sales}")

# テーブル5から機種構成データを抽出
machine_data = []
if len(doc.tables) > 5:
    table = doc.tables[5]
    for row_idx, row in enumerate(table.rows):
        if row_idx > 0 and row_idx < 7:
            row_data = [cell.text for cell in row.cells]
            machine_type = row_data[0]
            percentage = row_data[2]
            machine_data.append({
                'type': machine_type,
                'percentage': float(percentage.rstrip('%'))
            })
            print(f"機種構成: {machine_type} = {percentage}")

# ====================================
# グラフ 1: 売上シナリオ比較（棒グラフ）
# ====================================
fig1, ax1 = plt.subplots(figsize=(12, 7), dpi=150)

scenarios = [s['scenario'] for s in sales_data]
sales_values = []
for s in sales_data:
    sales_str = s['sales'].replace('約', '').replace('億円', '').strip()
    try:
        sales_values.append(float(sales_str))
    except:
        sales_values.append(0)

# 英語ラベルで生成
colors1 = ['#3498db', '#2ecc71', '#e74c3c']
x_pos = np.arange(len(scenarios))
bars = ax1.bar(x_pos, sales_values, color=colors1, alpha=0.85, 
               edgecolor='#2c3e50', linewidth=2.5, width=0.6)

# データラベル
for i, (bar, val) in enumerate(zip(bars, sales_values)):
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width()/2., height,
             f'{val:.1f}',
             ha='center', va='bottom', fontsize=13, fontweight='bold',
             bbox=dict(boxstyle='round,pad=0.5', facecolor='white', alpha=0.8))

ax1.set_ylabel('Annual Sales (100M Yen)', fontsize=13, fontweight='bold')
ax1.set_xlabel('Scenario', fontsize=13, fontweight='bold')
ax1.set_title('GiGO Abiko Sales Scenario Comparison', fontsize=15, fontweight='bold', pad=20)
ax1.set_xticks(x_pos)
ax1.set_xticklabels(['Conservative', 'Median', 'Model Store'], fontsize=11, fontweight='bold')
ax1.grid(axis='y', alpha=0.3, linestyle='--', linewidth=0.8)
ax1.set_ylim(0, max(sales_values) * 1.2)
ax1.tick_params(axis='y', labelsize=11)
ax1.set_facecolor('#f8f9fa')
fig1.patch.set_facecolor('white')

plt.tight_layout()
plt.savefig('sales_chart_temp.png', dpi=300, bbox_inches='tight', facecolor='white')
print("\n✓ Temporary sales chart created")
plt.close()

# ====================================
# グラフ 2: 機種構成（円グラフ）
# ====================================
fig2, ax2 = plt.subplots(figsize=(11, 9), dpi=150)

machine_types = [m['type'] for m in machine_data]
machine_percentages = [m['percentage'] for m in machine_data]

colors2 = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8']

wedges, texts, autotexts = ax2.pie(
    machine_percentages, 
    labels=[],  # 後でPILで追加
    autopct='%1.0f%%',
    colors=colors2, 
    startangle=90,
    textprops={'fontsize': 10, 'fontweight': 'bold'},
    explode=[0.08, 0.05, 0.05, 0.05, 0.05],
    shadow=True
)

for autotext in autotexts:
    autotext.set_color('white')
    autotext.set_fontsize(12)
    autotext.set_fontweight('bold')

ax2.set_title('Machine Type Distribution (460 units)', fontsize=15, fontweight='bold', pad=20)
fig2.patch.set_facecolor('white')

plt.tight_layout()
plt.savefig('machine_chart_temp.png', dpi=300, bbox_inches='tight', facecolor='white')
print("✓ Temporary machine chart created")
plt.close()

print("\n" + "="*60)
print("Temporary charts created - now adding Japanese labels...")
print("="*60)

# ====================================
# PILで日本語テキストを追加
# ====================================

# フォント候補
font_paths = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
]

# 英語フォント（フォールバック用）
default_font = None
try:
    default_font = ImageFont.load_default()
except:
    pass

# 1. 売上チャートに日本語タイトルを追加
try:
    img1 = Image.open('sales_chart_temp.png')
    draw1 = ImageDraw.Draw(img1)
    width, height = img1.size
    
    # 元のタイトル位置を置き換える（大概Y=60-80付近）
    # 黒いテキストでタイトルを描画
    title_y = 30
    title_text = "GiGO我孫子 売上シナリオ比較"
    
    # テキストサイズ調整
    if default_font:
        draw1.text((width/2 - 60, title_y), title_text, fill='black', font=default_font, anchor='mm')
    
    img1.save('sales_chart.png', quality=95)
    print("✓ sales_chart.png - Japanese labels added")
except Exception as e:
    print(f"⚠ Could not add Japanese labels to sales chart: {e}")
    # フォールバック：テンポラリファイルをコピー
    os.rename('sales_chart_temp.png', 'sales_chart.png')

# 2. 機種構成チャートに日本語ラベルを追加
try:
    img2 = Image.open('machine_chart_temp.png')
    draw2 = ImageDraw.Draw(img2)
    width, height = img2.size
    
    # 元のタイトル位置を置き換える
    title_y = 30
    title_text = "機種構成（460台ベース）"
    
    if default_font:
        draw2.text((width/2 - 60, title_y), title_text, fill='black', font=default_font, anchor='mm')
    
    img2.save('machine_chart.png', quality=95)
    print("✓ machine_chart.png - Japanese labels added")
except Exception as e:
    print(f"⚠ Could not add Japanese labels to machine chart: {e}")
    # フォールバック：テンポラリファイルをコピー
    os.rename('machine_chart_temp.png', 'machine_chart.png')

# テンポラリファイルを削除
for temp_file in ['sales_chart_temp.png', 'machine_chart_temp.png']:
    if os.path.exists(temp_file):
        os.remove(temp_file)

print("\n" + "="*60)
print("Complete: Presentation charts generated")
print("="*60)
print("  📊 sales_chart.png       - Sales scenario comparison")
print("  📊 machine_chart.png     - Machine type distribution")
print("="*60)
