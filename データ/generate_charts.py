#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
from docx import Document
import os

# Linux環境で日本語フォント設定（Noto Sans CJK JP を使用）
matplotlib.rcParams['font.sans-serif'] = ['Noto Sans CJK JP', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

# システムフォントを確認・追加
font_paths = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
]

for font_path in font_paths:
    if os.path.exists(font_path):
        try:
            fm.fontManager.addfont(font_path)
            print(f"✓ フォント追加: {font_path}")
            break
        except Exception as e:
            continue

# docxファイルから売上データを抽出
doc = Document('GiGO我孫子_買収計画_総合まとめ.docx')

# テーブル7から売上試算データを抽出（テーブル7＝インデックス7）
sales_data = []
if len(doc.tables) > 7:
    table = doc.tables[7]
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:  # ヘッダー
            headers = [cell.text for cell in row.cells]
        else:
            row_data = [cell.text for cell in row.cells]
            sale_scenario = row_data[0]  # シナリオ名
            annual_sales = row_data[3]   # 年間売上
            sales_data.append({'scenario': sale_scenario, 'sales': annual_sales})
            print(f"抽出: {sale_scenario} = {annual_sales}")

# テーブル5から機種構成データを抽出
machine_data = []
if len(doc.tables) > 5:
    table = doc.tables[5]
    for row_idx, row in enumerate(table.rows):
        if row_idx > 0 and row_idx < 7:  # クレーンのみのデータ行
            row_data = [cell.text for cell in row.cells]
            machine_type = row_data[0]
            machine_count = row_data[1]
            percentage = row_data[2]
            machine_data.append({
                'type': machine_type,
                'count': machine_count,
                'percentage': float(percentage.rstrip('%'))
            })
            print(f"抽出: {machine_type} = {percentage}")

# ====================================
# グラフ 1: 売上シナリオ比較（棒グラフ）
# ====================================
fig1, ax1 = plt.subplots(figsize=(12, 7))

scenarios = [s['scenario'] for s in sales_data]
# 年間売上から数値部分を抽出（例：「約2.5億円」→ 2.5）
sales_values = []
for s in sales_data:
    sales_str = s['sales'].replace('約', '').replace('億円', '').strip()
    try:
        sales_values.append(float(sales_str))
    except:
        sales_values.append(0)

colors1 = ['#3498db', '#2ecc71', '#e74c3c']
bars = ax1.bar(scenarios, sales_values, color=colors1, alpha=0.85, 
               edgecolor='#2c3e50', linewidth=2.5, width=0.6)

# データラベル
for bar, val in zip(bars, sales_values):
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width()/2., height,
             f'{val:.1f}億円',
             ha='center', va='bottom', fontsize=13, fontweight='bold',
             bbox=dict(boxstyle='round,pad=0.5', facecolor='white', alpha=0.8))

ax1.set_ylabel('年間売上', fontsize=13, fontweight='bold')
ax1.set_xlabel('シナリオ', fontsize=13, fontweight='bold')
ax1.set_title('GiGO我孫子 売上シナリオ比較', fontsize=15, fontweight='bold', pad=20)
ax1.grid(axis='y', alpha=0.3, linestyle='--', linewidth=0.8)
ax1.set_ylim(0, max(sales_values) * 1.2)
ax1.tick_params(axis='both', labelsize=11)

# 背景色
ax1.set_facecolor('#f8f9fa')
fig1.patch.set_facecolor('white')

plt.tight_layout()
plt.savefig('sales_chart.png', dpi=300, bbox_inches='tight', facecolor='white')
print("\n✓ sales_chart.png を生成しました")
plt.close()

# ====================================
# グラフ 2: 機種構成（円グラフ）
# ====================================
fig2, ax2 = plt.subplots(figsize=(11, 9))

machine_types = [m['type'] for m in machine_data]
machine_percentages = [m['percentage'] for m in machine_data]

# 色定義
colors2 = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8']

wedges, texts, autotexts = ax2.pie(
    machine_percentages, 
    labels=machine_types, 
    autopct='%1.0f%%',
    colors=colors2, 
    startangle=90,
    textprops={'fontsize': 11, 'fontweight': 'bold'},
    explode=[0.08, 0.05, 0.05, 0.05, 0.05],
    shadow=True
)

# テキスト装飾
for text in texts:
    text.set_fontsize(12)
    text.set_fontweight('bold')

for autotext in autotexts:
    autotext.set_color('white')
    autotext.set_fontsize(12)
    autotext.set_fontweight('bold')
    autotext.set_path_effects([matplotlib.patheffects.withStroke(linewidth=2, foreground='black')])

ax2.set_title('機種構成（460台ベース）', fontsize=15, fontweight='bold', pad=20)

# 背景色
fig2.patch.set_facecolor('white')

plt.tight_layout()
plt.savefig('machine_chart.png', dpi=300, bbox_inches='tight', facecolor='white')
print("✓ machine_chart.png を生成しました")
plt.close()

print("\n" + "="*60)
print("完了: プレゼン用グラフが生成されました")
print("="*60)
print(f"  📊 sales_chart.png       - 売上シナリオ比較")
print(f"  📊 machine_chart.png     - 機種構成（460台ベース）")
print("="*60)
