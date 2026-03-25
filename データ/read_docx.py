#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# docxファイルを読み込む
doc = Document('GiGO我孫子_買収計画_総合まとめ.docx')

print("=" * 80)
print("ドキュメント内容の確認")
print("=" * 80)

# テキストを抽出
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n" + "=" * 80)
print("テーブル内容の確認")
print("=" * 80)

# テーブルを確認
for table_idx, table in enumerate(doc.tables):
    print(f"\n【テーブル {table_idx}】")
    for row_idx, row in enumerate(table.rows):
        row_data = [cell.text for cell in row.cells]
        print(f"  行{row_idx}: {row_data}")
